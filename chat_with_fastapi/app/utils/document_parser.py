"""Turn an uploaded internal document into retrievable text chunks.

Two pure, side-effect-free steps, kept out of the service layer so they can be
unit-tested without a database, OpenAI key, or Chroma server:

  extract_text()  bytes  -> one plain-text string
  chunk_text()    string -> overlapping chunks sized for embedding
"""
import io
import re

# Extension -> the parser branch used for it. Anything not listed is rejected
# at upload time rather than silently indexed as mojibake.
SUPPORTED_EXTENSIONS = {
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
    ".pdf": "pdf",
    ".docx": "docx",
    ".html": "html",
    ".htm": "html",
}


class UnsupportedDocumentError(ValueError):
    """Raised for a file extension this parser has no branch for."""


def extension_of(filename: str) -> str:
    _, dot, ext = filename.rpartition(".")
    return f".{ext.lower()}" if dot else ""


def is_supported(filename: str) -> bool:
    return extension_of(filename) in SUPPORTED_EXTENSIONS


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded file's raw bytes.

    Dispatches on extension rather than the browser-supplied MIME type, which
    is unreliable (.md commonly arrives as application/octet-stream).
    """
    kind = SUPPORTED_EXTENSIONS.get(extension_of(filename))
    if kind is None:
        raise UnsupportedDocumentError(
            f"Unsupported file type '{extension_of(filename) or filename}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if kind == "text":
        return _normalize(data.decode("utf-8", errors="replace"))
    if kind == "pdf":
        return _normalize(_extract_pdf(data))
    if kind == "docx":
        return _normalize(_extract_docx(data))
    return _normalize(_extract_html(data))


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    # Blank-line separated so page boundaries survive as paragraph breaks and
    # chunk_text() can prefer splitting there.
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]

    # Tables hold a lot of the useful content in internal docs (policies,
    # limits, contacts) and are not part of `document.paragraphs`.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(part for part in parts if part.strip())


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def _normalize(text: str) -> str:
    """Collapse the whitespace noise extractors leave behind.

    Trailing spaces and runs of 3+ blank lines would otherwise eat into the
    chunk budget and blur the paragraph boundaries chunking relies on.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Split text into chunks of at most `chunk_size` characters.

    Splits on paragraph boundaries where possible so a chunk rarely cuts
    mid-sentence, then prefixes each chunk after the first with the tail of its
    predecessor. That overlap is what keeps an answer retrievable when it
    straddles a boundary — the cost is that a chunk can reach
    `chunk_size + overlap` characters.
    """
    text = _normalize(text)
    if not text:
        return []
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    overlap = max(0, min(overlap, chunk_size - 1))

    # Paragraphs first; any single paragraph longer than the budget is hard-split.
    pieces: list[str] = []
    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if len(paragraph) <= chunk_size:
            pieces.append(paragraph)
        else:
            pieces.extend(_split_oversized(paragraph, chunk_size))

    # Pack consecutive paragraphs together until the next one wouldn't fit.
    packed: list[str] = []
    current = ""
    for piece in pieces:
        if not current:
            current = piece
        elif len(current) + 2 + len(piece) <= chunk_size:
            current = f"{current}\n\n{piece}"
        else:
            packed.append(current)
            current = piece
    if current:
        packed.append(current)

    if overlap == 0:
        return packed

    return [
        chunk if i == 0 else f"{_tail(packed[i - 1], overlap)}\n\n{chunk}"
        for i, chunk in enumerate(packed)
    ]


def _split_oversized(paragraph: str, chunk_size: int) -> list[str]:
    """Break one over-long paragraph at whitespace near the size limit."""
    pieces: list[str] = []
    remaining = paragraph
    while len(remaining) > chunk_size:
        window = remaining[:chunk_size]
        cut = window.rfind(" ")
        # No whitespace in the whole window (minified text, CJK): hard-cut it.
        if cut <= chunk_size // 2:
            cut = chunk_size
        pieces.append(remaining[:cut].strip())
        remaining = remaining[cut:].strip()
    if remaining:
        pieces.append(remaining)
    return pieces


def _tail(text: str, size: int) -> str:
    """Last `size` characters, snapped forward to a word boundary."""
    if len(text) <= size:
        return text
    tail = text[-size:]
    space = tail.find(" ")
    return tail[space + 1:] if 0 <= space < size // 2 else tail
